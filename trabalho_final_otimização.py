#!pip install pulp

import pandas as pd
from pulp import LpProblem, LpMinimize, LpVariable, lpSum, LpStatus, value
from random import shuffle
from tabulate import tabulate
import easygui
import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from PIL import Image, ImageTk
from docx import Document



def ler_dados_excel(caminho_arquivo):
    df = pd.read_excel(caminho_arquivo)
    nutrient_columns = ['Calorias(kcal)', 'Lipídeos(g)', 'Proteína(g)', 'Potássio(mg)', 'Sódio(mg)',
                        'Fósforo(mg)', 'Cálcio(mg)', 'Ferro(mg)', 'Vitamina C(mg)', 'Carboidrato(g)']
    for col in nutrient_columns:
        df = df[pd.to_numeric(df[col], errors='coerce').notnull()]
    return df

def calcular_calorias(sexo, peso, altura, idade, nivel_atividade):
    if sexo.lower() == "homem":
        calorias = (10 * peso) + (6.25 * altura) - (5 * idade) + 5
    elif sexo.lower() == "mulher":
        calorias = (10 * peso) + (6.25 * altura) - (5 * idade) - 161
    else:
        print("Sexo inválido. Use 'homem' ou 'mulher'.")
        return None

    return calorias * nivel_atividade

def dividir_refeicoes(calorias_totais):
    refeicao1 = calorias_totais * 0.25
    refeicao2 = calorias_totais * 0.35
    refeicao3 = calorias_totais * 0.15
    refeicao4 = calorias_totais * 0.25
    return refeicao1, refeicao2, refeicao3, refeicao4

def criar_problema_otimizacao(df):
    prob = LpProblem("Dieta", LpMinimize)
    food_items = list(df['Descrição do alimento'])
    nutrientes = ['Calorias(kcal)', 'Lipídeos(g)', 'Proteína(g)', 'Potássio(mg)', 'Sódio(mg)',
                  'Fósforo(mg)', 'Cálcio(mg)', 'Ferro(mg)', 'Vitamina C(mg)', 'Carboidrato(g)']
    food_vars = LpVariable.dicts("Food", food_items, lowBound=0, cat='Continuous')

    # Adiciona restrições para limitar a quantidade de cada alimento
    max_quantity_per_food = 1
    for food in food_items:
        prob += food_vars[food] <= max_quantity_per_food

    # Adiciona restrições de nutrientes
    nutrient_limits = {'Potássio(mg)': 3500, 'Fósforo(mg)': 700, 'Proteína(g)': 60, 'Sódio(mg)': 2.3,
                       'Lipídeos(g)': 60, 'Cálcio(mg)': 1000, 'Ferro(mg)': 14, 'Vitamina C(mg)': 90,
                       'Carboidrato(g)': 300, 'Calorias(kcal)': 2000}
    for nutrient, limit in nutrient_limits.items():
        prob += lpSum([df.loc[df['Descrição do alimento'] == food, nutrient].values[0] * food_vars[food]
                       for food in food_items]) >= limit

    # Define a função objetivo
    prob += lpSum([df.loc[df['Descrição do alimento'] == food, 'Calorias(kcal)'].values[0] * food_vars[food]
                   for food in food_items])
    return prob, food_items

def resolver_problema_otimizacao(prob):
    status = prob.solve()
    if LpStatus[status] == 'Optimal':
        print('Solução ótima encontrada!')
        print("O valor do objetivo é:", value(prob.objective))
        for v in prob.variables():
            if v.varValue > 0:
              continue
    else:
        print('Solução não encontrada')

def distribuir_refeicoes(df, prob, food_items, num_refeicoes):
    refeicoes = {'Refeição 1': [], 'refeição 2': [], 'Refeição 3': [], 'Refeição 4': []}

    # Obtém a solução ótima do problema de otimização
    status = prob.solve()
    if LpStatus[status] != 'Optimal':
        print('Solução não encontrada')
        return refeicoes

    # Obtém os valores das variáveis de decisão
    food_values = {food.varValue: food for food in prob.variables()}

    # Ordena as variáveis de decisão com base nos valores
    sorted_food_values = sorted(food_values.items(), reverse=True)

    # Seleciona todos os alimentos da solução ótima
    selected_foods = [food for food in prob.variables() if food.varValue > 0]  # alimentos por refeição

    # Embaralha a lista de alimentos selecionados
    shuffle(selected_foods)

    # Divide os alimentos selecionados entre as refeições
    for meal in refeicoes.keys():
        refeicao = selected_foods[:len(selected_foods) // 4]  # Divide os alimentos em quatro partes
        refeicoes[meal].extend(refeicao)
        selected_foods = selected_foods[len(selected_foods) // 4:]  # Remove os alimentos já distribuídos
        #print(f'{meal}: {refeicao}')

    return refeicoes

# Ler dados do arquivo Excel
df = ler_dados_excel(r'C:\Users\RA00340509\Documents\VERSAO FINAL-20240424T113030Z-001\VERSAO FINAL\Base tratada VF.xlsx')


# Função para obter as necessidades diárias de calorias e dividir em refeições
def obter_necessidades_e_refeicoes():
    sexo = sexo_entry.get().lower()
    idade_str = idade_entry.get()
    peso_str = peso_entry.get()
    altura_str = altura_entry.get()
    nivel_atividade_str = nivel_atividade_entry.get()

    doc = Document()
    

    # Verificar se os campos estão vazios
    if idade_str == "" or peso_str == "" or altura_str == "" or nivel_atividade_str == "":
        messagebox.showerror("Erro", "Por favor, preencha todos os campos.")
        return
    if idade_str == "0" or peso_str == "0" or altura_str == "0" or nivel_atividade_str == "0":
        messagebox.showerror("Erro", "Por favor, insira valores diferentes de zero.")
        return

    # Converter os valores para inteiros ou floats
    try:
        idade = int(idade_str)
        peso = float(peso_str)
        altura = float(altura_str)
        nivel_atividade = float(nivel_atividade_str)
    except ValueError:
        messagebox.showerror("Erro", "Por favor, insira valores válidos.")
        return

    calorias_totais = calcular_calorias(sexo, peso, altura, idade, nivel_atividade)
    refeicao1, refeicao2, refeicao3, refeicao4 = dividir_refeicoes(calorias_totais)
    #doc.add_paragraph("Resultado")
    #doc.add_paragraph(f"Calorias totais: {calorias_totais}\nRefeição 1: {refeicao1}\nRefeição 2: {refeicao2}\nRefeição 3: {refeicao3}\nRefeição 4: {refeicao4}")
   
    # Obter as necessidades diárias de calorias e dividir em refeições
    #calorias_totais, refeicao1, refeicao2, refeicao3, refeicao4 = obter_necessidades_e_refeicoes()
    

    doc.add_paragraph(f"Sexo: {sexo}")
    doc.add_paragraph(f"Idade: {idade}")
    doc.add_paragraph(f"Peso: {peso}")
    doc.add_paragraph(f"Altura: {altura}")
    doc.add_paragraph(f"Nível de atividade: {nivel_atividade}")
    doc.add_paragraph(f"Calorias totais: {calorias_totais}")
    doc.add_paragraph(f"Refeição 1: {refeicao1}")
    doc.add_paragraph(f"Refeição 2: {refeicao2}")
    doc.add_paragraph(f"Refeição 3: {refeicao3}")
    doc.add_paragraph(f"Refeição 4: {refeicao4}")
    # Chame a função e atribua os valores retornados às variáveis
    #calorias_totais, refeicao1, refeicao2, refeicao3, refeicao4 = obter_necessidades_e_refeicoes()

    # Criar problema de otimização
    prob, food_items = criar_problema_otimizacao(df)

    # Resolver problema de otimização
    resolver_problema_otimizacao(prob)

    # Distribuir alimentos em refeições
    num_refeicoes = 4  # Defina o número de refeições que você deseja gerar
    refeicoes = distribuir_refeicoes(df, prob, food_items, num_refeicoes)

    doc.add_paragraph("\nAlimentos da Solução Ótima:")
    food_table = []
    for v in prob.variables():
        if v.varValue > 0:
            food_name = str(v).split('_')[1].replace(',', '')  # Obtém o nome do alimento
            food_table.append([food_name, v.varValue])  # Adiciona o nome e a quantidade à tabela

    doc.add_paragraph(tabulate(food_table, headers=['Alimento', 'Quantidade'], tablefmt='grid'))

    # Adicionando tabela de distribuição das calorias em refeições
    doc.add_paragraph("\nDistribuição das Calorias em Refeições:")
    calories_distribution_table = [
    ['Refeição 1', refeicao1],
    ['Refeição 2', refeicao2],
    ['Refeição 3', refeicao3],
    ['Refeição 4', refeicao4]
    ]
    #return calorias_totais, refeicao1, refeicao2, refeicao3, refeicao4
    doc.add_paragraph(tabulate(calories_distribution_table, headers=['Refeição', 'Calorias'], tablefmt='grid'))




    doc.add_paragraph("\nDistribuição das Refeições:")
    table = []
    for meal, food_vars in refeicoes.items():
        foods = [str(food).split('_')[1].replace(',', '') for food in food_vars]
        table.append([meal, ', '.join(foods)])
    doc.add_paragraph(tabulate(table, headers=['Refeições', 'Alimentos'], tablefmt='grid'))
    file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")])
    if file_path:
        doc.save(file_path)

    



    messagebox.showinfo("Sucesso", "Dieta salva com sucesso") 

    root.destroy() 
    return calorias_totais, refeicao1, refeicao2, refeicao3, refeicao4


    



def on_enter(e):
    enviar_lable.config(image=enviar_hover_image)

def on_leave(e):
    enviar_lable.config(image=enviar_image)

root = tk.Tk()
root.title("Obter Necessidades e Refeições")
root.geometry("1920x1080")

# Carregar imagens
background_image = Image.open("C:/Users/RA00340509/Documents/VERSAO FINAL-20240424T113030Z-001/VERSAO FINAL/Visual/f.png")
background_photo = ImageTk.PhotoImage(background_image)

global enviar_photo, enviar_hover_photo  # Add this line
enviar_image = Image.open("C:/Users/RA00340509/Documents/VERSAO FINAL-20240424T113030Z-001/VERSAO FINAL/Visual/e.png")
enviar_photo = ImageTk.PhotoImage(enviar_image)

enviar_hover_image = Image.open("C:/Users/RA00340509/Documents/VERSAO FINAL-20240424T113030Z-001/VERSAO FINAL/Visual/eselecionado.png")
enviar_hover_photo = ImageTk.PhotoImage(enviar_hover_image)

# Configurar o fundo
background_label = tk.Label(root, image=background_photo)  # Use background_photo instead of root.background_photo
background_label.place(relwidth=1, relheight=1)



# Caixas de entrada
sexo_label = ttk.Label(root, text="Informe o sexo (mulher ou homem):")
sexo_label.place(x=700, y=430)
sexo_entry = ttk.Entry(root)
sexo_entry.place(x=900, y=430)


idade_label = ttk.Label(root, text="Informe a idade:")
idade_label.place(x=700, y=460)
idade_entry = ttk.Entry(root)
idade_entry.place(x=900, y=460)

peso_label = ttk.Label(root, text="Informe o peso em kg:")
peso_label.place(x=700, y=490)
peso_entry = ttk.Entry(root)
peso_entry.place(x=900, y=490)

altura_label = ttk.Label(root, text="Informe a altura em cm:")
altura_label.place(x=700, y=520)
altura_entry = ttk.Entry(root)
altura_entry.place(x=900, y=520)

nivel_atividade_label = ttk.Label(root, text="Informe o nível de atividade física (1 para leve, 2 para moderado, 3 para elevado ou 4 para intenso):")
nivel_atividade_label.place(x=700, y=550)
nivel_atividade_entry = ttk.Entry(root)
nivel_atividade_entry.place(x=900, y=570)

# Botão de enviar
enviar_lable = tk.Button(root, image=enviar_photo, borderwidth=0, command=obter_necessidades_e_refeicoes)
enviar_lable.bind("<Enter>", lambda e: enviar_lable.config(image=enviar_hover_photo))
enviar_lable.bind("<Leave>", lambda e: enviar_lable.config(image=enviar_photo))
enviar_lable.pack()
enviar_lable.place(x=650, y=700)  # Posiciona o botão nas coordenadas (50, 50)



# Label para mostrar os resultados
resultado_label = ttk.Label(root, text="")
resultado_label.pack()


root.mainloop()

